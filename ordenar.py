import re
import docx
import os
from docx.enum.text import WD_ALIGN_PARAGRAPH 
from docx.shared import Inches 
from spire.doc.common import *
from spire.doc import *

naoIndentificado = []
equalizador = []
fonte12v8a = []
fonte24v6a = []
fonte48v15a = []
fonte48v30a = []
fonte48v40a = []
inversor1000w = []
inversor3000w = []
inversor5000w = []

def read_docx(file_path):
    doc = docx.Document(file_path)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return text

def read_text(text):
    items = []
    current_item = {}

    # Divide o texto em itens separados
    item_texts = re.split(r'-{5,}', text)

    for item_text in item_texts:
        lines = item_text.strip().split('\n')
        for line in lines:
            if line.startswith("Modelo:"):
                if current_item:
                    process_item(current_item)
                    current_item = {}
                current_item['Modelo'] = line.split("Modelo:", 1)[1].strip()
            elif line.startswith("URL:"):
                current_item['URL'] = line.split("URL:", 1)[1].strip()
            elif line.startswith("Nome:"):
                current_item['Nome'] = line.split("Nome:", 1)[1].strip()
            elif line.startswith("Preço:"):
                current_item['Preço'] = line.split("Preço:", 1)[1].strip()
            elif line.startswith("Preço Previsto:"):
                current_item['Preço Previsto'] = line.split("Preço Previsto:", 1)[1].strip()
            elif line.startswith("Loja:"):
                current_item['Loja'] = line.split("Loja:", 1)[1].strip()
            elif line.startswith("Tipo:"):
                current_item['Tipo'] = line.split("Tipo:", 1)[1].strip()
            elif line.startswith("Lugar:"):
                current_item['Lugar'] = line.split("Lugar:", 1)[1].strip()
            elif line.startswith("Cupom:"):
                current_item['Cupom'] = line.split("Cupom:", 1)[1].strip()
                
        if current_item:
            process_item(current_item)
            current_item = {}

def process_item(item):
    if item['Modelo'] == "Nao indentificado":
        naoIndentificado.append(format_item_dif(item))
    elif item['Modelo'] == "EQUALIZADOR PARA BANCO DE BATERIAS":
        equalizador.append((format_item(item), item['Loja']))
    elif item['Modelo'] == "FONTE NOBREAK 12V/8A":
        fonte12v8a.append((format_item(item), item['Loja']))
    elif item['Modelo'] == "FONTE NOBREAK 24V/6A":
        fonte24v6a.append((format_item(item), item['Loja']))
    elif item['Modelo'] == "FONTE NOBREAK -48V 15A 15A":
        fonte48v15a.append((format_item(item), item['Loja']))
    elif item['Modelo'] == "FONTE NOBREAK -48V 30A 15A":
        fonte48v30a.append((format_item(item), item['Loja']))
    elif item['Modelo'] == "FONTE NOBREAK -48V 40A 10A":
        fonte48v40a.append((format_item(item), item['Loja']))
    elif item['Modelo'] == "INVERSOR OFF GRID SENOIDAL PURA JFA 1000W 48V/220V RACK":
        inversor1000w.append((format_item(item), item['Loja']))
    elif item['Modelo'] == "INVERSOR OFF GRID SENOIDAL PURA JFA 3000W 48/220V C/ GER RACK":
        inversor3000w.append((format_item(item), item['Loja']))
    elif item['Modelo'] == "INVERSOR OFF GRID SENOIDAL PURA JFA 5000W 48/220V C/ GER RACK":
        inversor5000w.append((format_item(item), item['Loja']))

def format_item(item):
    formatted_item = f"{item['Cupom']}  {item['Loja']} – {item['Lugar']} – PreçoAnúncio: R$ {item['Preço']} – Preço Política: R$ {item['Preço Previsto']} ({item['Tipo']}) {item['URL']}\n"
    return formatted_item

def format_item_dif(item):
    formatted_item = f"{item['URL']}\n"
    return formatted_item

lojas = {}

for item_path in os.listdir(r"dados/"):
    file_path = os.path.join(r"dados/", item_path)
    text = read_docx(file_path)
    read_text(text)

output_doc = docx.Document()

produtos = [
    (equalizador, "EQUALIZADOR PARA BANCO DE BATERIAS"),
    (fonte12v8a, "FONTE NOBREAK 12V/8A"),
    (fonte24v6a, "FONTE NOBREAK 24V/6A"),
    (fonte48v15a, "FONTE NOBREAK -48V 15A"),
    (fonte48v30a, "FONTE NOBREAK -48V 30A"),
    (fonte48v40a, "FONTE NOBREAK -48V 40A 10A"),
    (inversor1000w, "INVERSOR OFF GRID SENOIDAL PURA JFA 1000W 48V/220V RACK"),
    (inversor3000w, "INVERSOR OFF GRID SENOIDAL PURA JFA 3000W 48/220V C/ GER RACK"),
    (inversor5000w, "INVERSOR OFF GRID SENOIDAL PURA JFA 5000W 48/220V C/ GER RACK")
]

for lista, modelo in produtos:
    for item in lista:
        if item[1] not in lojas:
            lojas[item[1]] = []
        lojas[item[1]].append((item[0], modelo))

for i in lojas:
    output_doc.add_paragraph().add_run(f"*{i}*").bold = True
    for item, modelo in lojas[i]:
        output_doc.add_paragraph(f"{modelo} - {item}").paragraph_format.left_indent = Inches(0.5)
        
output_doc.save(r'dados_extraidos.docx')