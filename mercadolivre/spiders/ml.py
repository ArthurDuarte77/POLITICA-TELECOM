#IMPORTANTE
#//div/div/div[2]/div[2]/div[1]/div[2]/div[1]/p[@class="ui-promotions-pill ui-pb-highlight-wrapper coupon"]
import re
import requests
import unidecode
import scrapy
import requests
from docx import Document
import pandas
from datetime import datetime
import time

    
df = pandas.read_excel("PoliticadePrecos_Telecom.xlsx", engine='openpyxl', sheet_name="Sugestão de Preços", skiprows=1)
# equalizadorClassico = 0
# equalizadorPremium = 0
# fonte12v8aClassico = 0
# fonte12v8aPremium = 0
# fonte24v6aClassico = 0
# fonte24v6aPremium = 0
# fonte48v15aClassico = 0
# fonte48v15aPremium = 0
# fonte48v30aClassico = 0
# fonte48v30aPremium = 0
# fonte48v40aClassico = 0
# fonte48v40aPremium = 0
# inversor1000wClassico = 0
# inversor1000wPremium = 0
# inversor3000wClassico = 0
# inversor3000wPremium = 0
# inversor5000wClassico = 0
# inversor5000wPremium = 0
urls_visitadas = []

for index, i in df.iterrows():
    if i['CONSTRUÇÃO DE CUSTO '] == "EQUALIZADOR PARA BANCO DE BATERIAS":
        if i['Anuncio'] == "Classico":
            equalizadorClassico = round(i['PREÇO ML '], 2) - 0.01
        elif i['Anuncio'] == "Premium":
            equalizadorPremium = round(i['PREÇO ML '], 2) - 0.01
    elif i['CONSTRUÇÃO DE CUSTO '] == "FONTE NOBREAK 12V/8A":
        if i['Anuncio'] == "Classico":
            fonte12v8aClassico = round(i['PREÇO ML '], 2) - 0.01
        elif i['Anuncio'] == 'Premium':
            fonte12v8aPremium = round(i['PREÇO ML '], 2) - 0.01
        
    elif i['CONSTRUÇÃO DE CUSTO '] == "FONTE NOBREAK 24V/6A":
        if i['Anuncio'] == "Classico":
            fonte24v6aClassico = round(i['PREÇO ML '], 2) - 0.01
        elif i['Anuncio'] == 'Premium':
            fonte24v6aPremium = round(i['PREÇO ML '], 2) - 0.01
        
    elif i['CONSTRUÇÃO DE CUSTO '] == "FONTE NOBREAK -48V 15A 15A":
        if i['Anuncio'] == "Classico":
            fonte48v15aClassico = round(i['PREÇO ML '], 2) - 0.01
        elif i['Anuncio'] == 'Premium':
            fonte48v15aPremium = round(i['PREÇO ML '], 2) - 0.01
        
    elif i['CONSTRUÇÃO DE CUSTO '] == "FONTE NOBREAK -48V 30A 15A":
        if i['Anuncio'] == "Classico":
            fonte48v30aClassico = round(i['PREÇO ML '], 2) - 0.01
        elif i['Anuncio'] == 'Premium':
            fonte48v30aPremium = round(i['PREÇO ML '], 2) - 0.01
        
    elif i['CONSTRUÇÃO DE CUSTO '] == "FONTE NOBREAK -48V 40A 10A":
        if i['Anuncio'] == "Classico":
            fonte48v40aClassico = round(i['PREÇO ML '], 2) - 0.01
        elif i['Anuncio'] == 'Premium':
            fonte48v40aPremium = round(i['PREÇO ML '], 2) - 0.01
        
    elif i['CONSTRUÇÃO DE CUSTO '] == "INVERSOR OFF GRID SENOIDAL PURA JFA 1000W 48V/220V  RACK ":
        if i['Anuncio'] == "Classico":
            inversor1000wClassico = round(i['PREÇO ML '], 2) - 0.01
        elif i['Anuncio'] == 'Premium':
            inversor1000wPremium = round(i['PREÇO ML '], 2) - 0.01
        
    elif i['CONSTRUÇÃO DE CUSTO '] == "INVERSOR OFF GRID SENOIDAL PURA JFA 3000W 48/220V C/  GER RACK ":
        if i['Anuncio'] == "Classico":
            inversor3000wClassico = round(i['PREÇO ML '], 2) - 0.01
        elif i['Anuncio'] == 'Premium':
            inversor3000wPremium = round(i['PREÇO ML '], 2) - 0.01
        
    elif i['CONSTRUÇÃO DE CUSTO '] == "INVERSOR OFF GRID SENOIDAL PURA JFA 5000W 48/220V C/  GER RACK ":
        if i['Anuncio'] == "Classico":
            inversor5000wClassico = round(i['PREÇO ML '], 2) - 0.01
        elif i['Anuncio'] == 'Premium':
            inversor5000wPremium = round(i['PREÇO ML '], 2) - 0.01
        
# if os.path.exists("dados_scrapy.docx"):
#     doc = Document("dados_scrapy.docx")
# else:
doc = Document()

def extract_price(response):
  price_selectors = [
      '//*[@id="price"]/div/div[1]/div[1]/span[1]/span/span[2]/text()',
      '//html/body/main/div[2]/div[5]/div/div[1]/div/div[1]/div/div[@class="ui-pdp-container__row ui-pdp-container__row--price"]/div/div[1]/div[1]/span/span/span[2]/text()',
      '//*[@id="ui-pdp-main-container"]/div[1]/div/div[1]/div[2]/div[3]/div[1]/div[1]/span/span/span[2]/text()',
      '//*[@id="ui-pdp-main-container"]/div[1]/div/div[1]/div[2]/div[2]/div[1]/div[1]/span[1]/span/span[2]/text()'
  ]
  
  for selector in price_selectors:
    price = response.xpath(selector).get()
    if price:
      price = price.replace('.', '')
      decimal_selector = selector.replace("span[2]/text()", "") + 'span[@class="andes-money-amount__cents andes-money-amount__cents--superscript-36"]/text()'
      price_decimal = response.xpath(decimal_selector).get()
      
      if price_decimal:
        return float(f"{price}.{price_decimal}")
      else:
        try:
          return float(price)
        except ValueError:
          pass

  return None  


def extract_price_new(response):
  price_selectors = [
      './/div/div/div[2]/div[2]/div[1]/div[2]/div/div/div/span[1]/span[@class="andes-money-amount__fraction"]/text()',
      './/div/div/div[2]/div/div[2]/div/div/div/span[1]/span[@class="andes-money-amount__fraction"]/text()',
      './/div/div/div[2]/div/div[3]/div/div/div/span[1]/span[@class="andes-money-amount__fraction"]/text()',
      './/div[1]/div[1]/div/div/div/span[1]/span[@class="andes-money-amount__fraction"]/text()'
  ]
  
  for selector in price_selectors:
    price = response.xpath(selector).get()
    if price:
      price = price.replace('.', '')
      decimal_selector = selector.replace('span[@class="andes-money-amount__fraction"]/text()', '') + 'span[@class="andes-money-amount__cents andes-money-amount__cents--superscript-24"]/text()'
      price_decimal = response.xpath(decimal_selector).get()
      
      if price_decimal:
        return float(f"{price}.{price_decimal}")
      else:
        try:
          return float(price)
        except ValueError:
          pass

  return None  


class MlSpider(scrapy.Spider):
    option_selected = ""
    option_selected_new = ""

    name = 'ml'
    start_urls = ["https://lista.mercadolivre.com.br/fonte-jfa"]
    
    def __init__(self, palavra=None, cookie=None, *args, **kwargs):
        super(MlSpider, self).__init__(*args, **kwargs)
        self.palavra = palavra
        self.cookie = cookie
    
    
    def parse(self, response, **kwargs):
        self.option_selected = self.palavra
        self.option_selected_new = self.palavra
        search = ""
        if self.option_selected == "EQUALIZADOR PARA BANCO DE BATERIAS":
            search = "equalizador de baterias"
        elif self.option_selected == "FONTE NOBREAK 12V/8A":
            search = "fonte nobreak 12v 8a"
        elif self.option_selected == "FONTE NOBREAK 24V/6A":
            search = "fonte nobreak 24v 6a"
        elif self.option_selected == "FONTE NOBREAK -48V 15A 15A":
            search = "fonte nobreak -48v 15a 15a"
        elif self.option_selected == "FONTE NOBREAK -48V 30A 15A":
            search = "fonte nobreak -48v 30a 15a"
        elif self.option_selected == "FONTE NOBREAK -48V 40A 10A":
            search = "fonte nobreak -48v 40a 10a"
        elif self.option_selected == "INVERSOR OFF GRID SENOIDAL PURA JFA 1000W 48V/220V  RACK":
            search = "inversor 1000w 48/220v onda senoidal pura jfa ger rack"
        elif self.option_selected == "INVERSOR OFF GRID SENOIDAL PURA JFA 3000W 48/220V C/ GER RACK":
            search = "inversor 3000w 48/220v onda senoidal pura jfa ger rack"
        elif self.option_selected == "INVERSOR OFF GRID SENOIDAL PURA JFA 5000W 48/220V C/ GER RACK":
            search = "inversor 5000w 48/220v onda senoidal pura jfa ger rack"
        search = search.replace(" ", "-")
        # yield scrapy.Request(dont_filter=True, url=f"https://lista.mercadolivre.com.br/acessorios-veiculos/{search}_OrderId_PRICE_NoIndex_True", callback=self.parse_all)BRAND_22292586
        # yield scrapy.Request(dont_filter=True, url=f"https://lista.mercadolivre.com.br/acessorios-veiculos/{search}_Frete_Full_OrderId_PRICE_NoIndex_True", callback=self.parse_all)
        if self.option_selected == "EQUALIZADOR PARA BANCO DE BATERIAS":
            yield scrapy.Request(dont_filter=True, url=f"https://lista.mercadolivre.com.br/{search}_Frete_Full_OrderId_PRICE_NoIndex_True", callback=self.parse_all)
            yield scrapy.Request(dont_filter=True, url=f"https://lista.mercadolivre.com.br/{search}_OrderId_PRICE_NoIndex_True", callback=self.parse_all)
        else:
            yield scrapy.Request(dont_filter=True, url=f"https://lista.mercadolivre.com.br/{search}_Frete_Full_OrderId_PRICE_BRAND_2466336_NoIndex_True", callback=self.parse_all)
            yield scrapy.Request(dont_filter=True, url=f"https://lista.mercadolivre.com.br/{search}_OrderId_PRICE_BRAND_2466336_NoIndex_True", callback=self.parse_all)
            yield scrapy.Request(dont_filter=True, url=f"https://lista.mercadolivre.com.br/{search}_Frete_Full_OrderId_PRICE_BRAND_22292586_NoIndex_True", callback=self.parse_all)
            yield scrapy.Request(dont_filter=True, url=f"https://lista.mercadolivre.com.br/{search}_OrderId_PRICE_BRAND_22292586_NoIndex_True", callback=self.parse_all)
        
    
    def parse_all(self, response):
        for item in response.xpath('//div/div[3]/section/ol/li[@class="ui-search-layout__item"]'):
            # new_name = item.xpath('.//h2[@class="poly-box poly-component__title"]').get()
            new_name = item.xpath('.//h2[@class="ui-search-item__title"]/text()').get()
            if not new_name:
                new_name = item.xpath('.//h2[@class="ui-search-item__title ui-search-item__group__element"]/a/text()').get()
            name = new_name
            if not new_name:
                print(response.url)
            price = extract_price_new(response=item)
            if not price:
                print(response.url)
            cupom = ""
            # if item.xpath('.//div/div/div[2]/div[2]/div[1]/div[2]/div[1]/p[@class="ui-promotions-pill ui-pb-highlight-wrapper coupon"]').get():
            #     cupom = item.xpath('.//div/div/div[2]/div[2]/div[1]/div[2]/div[1]/p[@class="ui-promotions-pill ui-pb-highlight-wrapper coupon"]/span/span/span/text()').get().replace("OFF", "")
            #     if "%" in cupom and price:
            #         cupom = int(re.findall(r'(\d+)%', cupom)[0])
            #         cupom = f"Cupom: %{cupom} - {round(price - (price *( cupom / 100)), 2)}"
            #     elif "R$" in cupom and price:
            #         cupom = int(re.findall(r'R\$\s?(\d+,\d+|\d+)', cupom)[0])
            #         cupom = f"Cupom: R${cupom} - {round(price - cupom, 2)}"
            # if item.xpath('.//ul[@class="ui-search-winner-alternatives ui-search-winner-alternatives__container--top-space"]/li[@class="ui-search-winner-alternatives__item"]/a/div').get():
            #     cupom += " Mais de um item"
            loja = ""
            listing_type = "Not Found"
            if item.xpath('.//span[@class="ui-search-item__group__element ui-search-installments ui-search-color--BLACK"]').get():
                listing_type = "Clássico"
            elif item.xpath('.//span[@class="ui-search-item__group__element ui-search-installments ui-search-color--LIGHT_GREEN"]').get():
                listing_type = "Premium"
            url = item.xpath('.//h2[@class="poly-box poly-component__title"]/a/@href').get()
            new_name = unidecode.unidecode(new_name.lower())
            url = item.xpath('.//div/div/div[2]/div[1]/a[@class="ui-search-item__group__element ui-search-link__title-card ui-search-link"]/@href').get()
            if not url:
                url = item.xpath('.//a[@class="ui-search-item__group__element ui-search-link__title-card ui-search-link"]/@href').get()
            if not url:
                url = item.xpath('.//h2[@class="ui-search-item__title ui-search-item__group__element"]/a/@href').get()
            if "taramps" in new_name or "stetson" in new_name or "usina" in new_name:
                continue
            if self.option_selected == "EQUALIZADOR PARA BANCO DE BATERIAS":     
                if "equalizador" in new_name and ("bateria" in new_name or "baterias" in new_name) and "jfa" in new_name:
                    if "kit" in new_name:
                        kit_match = re.search(r'kit\s+(\d+)', new_name)
                        if kit_match:
                            num_kits = int(kit_match.group(1))
                            if num_kits > 1 and price:
                                price = round(price / num_kits, 2)
                                cupom = f"KIT: {num_kits} UNIDADES"
                    if listing_type == "Clássico" and price:
                        if price >= equalizadorClassico:
                            continue
                    elif listing_type == "Premium" and price:
                        if price >= equalizadorPremium:
                            continue
                    yield scrapy.Request(dont_filter=True, url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type, 'cupom': cupom})
                        
            elif self.option_selected == "FONTE NOBREAK 12V/8A":
                if "fonte" in new_name and "nobreak" in new_name and "12v" in new_name and "8a" in new_name:
                    if listing_type == "Clássico" and price and cupom == "":
                        if price >= fonte12v8aClassico:
                            continue
                    elif listing_type == "Premium" and price and cupom == "":
                        if price >= fonte12v8aPremium:
                            continue
                    yield scrapy.Request(dont_filter=True, url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type, 'cupom': cupom})

            elif self.option_selected == "FONTE NOBREAK 24V/6A":
                if "fonte" in new_name and "nobreak" in new_name and "24v" in new_name and "6a" in new_name:
                    if listing_type == "Clássico" and price and cupom == "":
                        if price >= fonte24v6aClassico:
                            continue
                    elif listing_type == "Premium" and price and cupom == "":
                        if price >= fonte24v6aPremium:
                            continue
                    yield scrapy.Request(dont_filter=True, url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type, 'cupom': cupom})

            elif self.option_selected == "FONTE NOBREAK -48V 15A 15A":
                if "fonte" in new_name and "nobreak" in new_name and "48v" in new_name and "15a" in new_name:
                    if listing_type == "Clássico" and price and cupom == "":
                        if price >= fonte48v15aClassico:
                            continue
                    elif listing_type == "Premium" and price and cupom == "":
                        if price >= fonte48v15aPremium:
                            continue
                    yield scrapy.Request(dont_filter=True, url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type, 'cupom': cupom})

            elif self.option_selected == "FONTE NOBREAK -48V 30A 15A":
                if "fonte" in new_name and "nobreak" in new_name and "48v" in new_name and "30a" in new_name:
                    if listing_type == "Clássico" and price and cupom == "":
                        if price >= fonte48v30aClassico:
                            continue
                    elif listing_type == "Premium" and price and cupom == "":
                        if price >= fonte48v30aPremium:
                            continue
                    yield scrapy.Request(dont_filter=True, url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type, 'cupom': cupom})

            elif self.option_selected == "FONTE NOBREAK -48V 40A 10A":
                if "fonte" in new_name and "nobreak" in new_name and ("48v" in new_name or "48" in new_name) and ("40a" in new_name or "40" in new_name):
                    if listing_type == "Clássico" and price and cupom == "":
                        if price >= fonte48v40aClassico:
                            continue
                    elif listing_type == "Premium" and price and cupom == "":
                        if price >= fonte48v40aPremium:
                            continue
                    yield scrapy.Request(dont_filter=True, url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type, 'cupom': cupom})

            elif self.option_selected == "INVERSOR OFF GRID SENOIDAL PURA JFA 1000W 48V/220V  RACK":
                if "inversor" in new_name and "senoidal" in new_name and "1000w" in new_name and ("48v" in new_name or "48" in new_name) and ("220v" in new_name or "220" in new_name) and "ger" in new_name and "rack" in new_name:
                    if listing_type == "Clássico" and price and cupom == "":
                        if price >= inversor1000wClassico:
                            continue
                    elif listing_type == "Premium" and price and cupom == "":
                        if price >= inversor1000wPremium:
                            continue
                    yield scrapy.Request(dont_filter=True, url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type, 'cupom': cupom})

            elif self.option_selected == "INVERSOR OFF GRID SENOIDAL PURA JFA 3000W 48/220V C/ GER RACK":
                if "inversor" in new_name and "senoidal" in new_name and "3000w" in new_name and ("48v" in new_name or "48" in new_name) and ("220v" in new_name or "220" in new_name) and "ger" in new_name and "rack" in new_name:
                    if listing_type == "Clássico" and price and cupom == "":
                        if price >= inversor3000wClassico:
                            continue
                    elif listing_type == "Premium" and price and cupom == "":
                        if price >= inversor3000wPremium:
                            continue
                    yield scrapy.Request(dont_filter=True, url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type, 'cupom': cupom})

            elif self.option_selected == "INVERSOR OFF GRID SENOIDAL PURA JFA 5000W 48/220V C/ GER RACK":
                if "inversor" in new_name and "senoidal" in new_name and "5000w" in new_name and ("48v" in new_name or "48" in new_name) and ("220v" in new_name or "220" in new_name) and "ger" in new_name and "rack" in new_name:
                    if listing_type == "Clássico" and price and cupom == "":
                        if price >= inversor5000wClassico:
                            continue
                    elif listing_type == "Premium" and price and cupom == "":
                        if price >= inversor5000wPremium:
                            continue
                    yield scrapy.Request(dont_filter=True, url=url, callback=self.parse_product, meta={'name': name, 'loja': loja, 'price':price, 'listing_type': listing_type, 'cupom': cupom})

        if response.xpath('//nav/ul/li/a[@class="andes-pagination__link" and @title="Seguinte"]'):
            next_page = response.xpath('//nav/ul/li/a[@class="andes-pagination__link" and @title="Seguinte"]/@href').get()
            if next_page:
                yield scrapy.Request(dont_filter=True, url=next_page, callback=self.parse_all)


                
    def parse_product(self, response):
        cupom = response.meta['cupom']
        name = response.meta['name']
        loja = response.meta['loja']
        listing_type = response.meta['listing_type']
        loja = response.xpath('//div[1]/div/button[@class="ui-pdp-seller__link-trigger-button non-selectable"]/span[2]/text()').get()
        self.option_selected_new = self.option_selected
        new_price_float = response.meta["price"]
        tipo = listing_type
        location_url = f'https://www.mercadolivre.com.br/perfil/{loja.replace(" ", "+")}'
        yield scrapy.Request(dont_filter=True, url=location_url, callback=self.parse_location, meta={'url': response.url, 'name': name, 'price': new_price_float, 'qtde_parcelado': 0, 'price_parcelado': 0, 'loja': loja, 'tipo': tipo, 'cupom': cupom})


    def get_price_previsto(self, tipo):
        if tipo == "Clássico":
            for index, i in df.iterrows():
                if self.option_selected_new == "EQUALIZADOR PARA BANCO DE BATERIAS" and i['CONSTRUÇÃO DE CUSTO '] == "EQUALIZADOR PARA BANCO DE BATERIAS" and i['Anuncio'] == "Classico":
                    return round(i['PREÇO ML '], 2)
                elif self.option_selected_new == "FONTE NOBREAK 12V/8A" and i['CONSTRUÇÃO DE CUSTO '] == "FONTE NOBREAK 12V/8A" and i['Anuncio'] == "Classico":
                    return round(i['PREÇO ML '], 2)
                elif self.option_selected_new == "FONTE NOBREAK 24V/6A" and i['CONSTRUÇÃO DE CUSTO '] == "FONTE NOBREAK 24V/6A" and i['Anuncio'] == "Classico":
                    return round(i['PREÇO ML '], 2)
                elif self.option_selected_new == "FONTE NOBREAK -48V 15A 15A" and i['CONSTRUÇÃO DE CUSTO '] == "FONTE NOBREAK -48V 15A 15A" and i['Anuncio'] == "Classico":
                    return round(i['PREÇO ML '], 2)
                elif self.option_selected_new == "FONTE NOBREAK -48V 30A 15A" and i['CONSTRUÇÃO DE CUSTO '] == "FONTE NOBREAK -48V 30A 15A" and i['Anuncio'] == "Classico":
                    return round(i['PREÇO ML '], 2)
                elif self.option_selected_new == "FONTE NOBREAK -48V 40A 10A" and i['CONSTRUÇÃO DE CUSTO '] == "FONTE NOBREAK -48V 40A 10A" and i['Anuncio'] == "Classico":
                    return round(i['PREÇO ML '], 2)
                elif self.option_selected_new == "INVERSOR OFF GRID SENOIDAL PURA JFA 1000W 48V/220V  RACK" and i['CONSTRUÇÃO DE CUSTO '] == "INVERSOR OFF GRID SENOIDAL PURA JFA 1000W 48V/220V  RACK " and i['Anuncio'] == "Classico":
                    return round(i['PREÇO ML '], 2)
                elif self.option_selected_new == "INVERSOR OFF GRID SENOIDAL PURA JFA 3000W 48/220V C/ GER RACK" and i['CONSTRUÇÃO DE CUSTO '] == "INVERSOR OFF GRID SENOIDAL PURA JFA 3000W 48/220V C/  GER RACK " and i['Anuncio'] == "Classico":
                    return round(i['PREÇO ML '], 2)
                elif self.option_selected_new == "INVERSOR OFF GRID SENOIDAL PURA JFA 5000W 48/220V C/ GER RACK " and i['CONSTRUÇÃO DE CUSTO '] == "INVERSOR OFF GRID SENOIDAL PURA JFA 5000W 48/220V C/  GER RACK " and i['Anuncio'] == "Classico":
                    return round(i['PREÇO ML '], 2)
        elif tipo == "Premium":
            for index, i in df.iterrows():
                if self.option_selected_new == "EQUALIZADOR PARA BANCO DE BATERIAS" and i['CONSTRUÇÃO DE CUSTO '] == "EQUALIZADOR PARA BANCO DE BATERIAS" and i['Anuncio'] == "Premium":
                    return round(i['PREÇO ML '], 2)
                elif self.option_selected_new == "FONTE NOBREAK 12V/8A" and i['CONSTRUÇÃO DE CUSTO '] == "FONTE NOBREAK 12V/8A" and i['Anuncio'] == "Premium":
                    return round(i['PREÇO ML '], 2)
                elif self.option_selected_new == "FONTE NOBREAK 24V/6A" and i['CONSTRUÇÃO DE CUSTO '] == "FONTE NOBREAK 24V/6A" and i['Anuncio'] == "Premium":
                    return round(i['PREÇO ML '], 2)
                elif self.option_selected_new == "FONTE NOBREAK -48V 15A 15A" and i['CONSTRUÇÃO DE CUSTO '] == "FONTE NOBREAK -48V 15A 15A" and i['Anuncio'] == "Premium":
                    return round(i['PREÇO ML '], 2)
                elif self.option_selected_new == "FONTE NOBREAK -48V 30A 15A" and i['CONSTRUÇÃO DE CUSTO '] == "FONTE NOBREAK -48V 30A 15A" and i['Anuncio'] == "Premium":
                    return round(i['PREÇO ML '], 2)
                elif self.option_selected_new == "FONTE NOBREAK -48V 40A 10A" and i['CONSTRUÇÃO DE CUSTO '] == "FONTE NOBREAK -48V 40A 10A" and i['Anuncio'] == "Premium":
                    return round(i['PREÇO ML '], 2)
                elif self.option_selected_new == "INVERSOR OFF GRID SENOIDAL PURA JFA 1000W 48V/220V  RACK" and i['CONSTRUÇÃO DE CUSTO '] == "INVERSOR OFF GRID SENOIDAL PURA JFA 1000W 48V/220V  RACK " and i['Anuncio'] == "Premium":
                    return round(i['PREÇO ML '], 2)
                elif self.option_selected_new == "INVERSOR OFF GRID SENOIDAL PURA JFA 3000W 48/220V C/ GER RACK" and i['CONSTRUÇÃO DE CUSTO '] == "INVERSOR OFF GRID SENOIDAL PURA JFA 3000W 48/220V C/  GER RACK " and i['Anuncio'] == "Premium":
                    return round(i['PREÇO ML '], 2)
                elif self.option_selected_new == "INVERSOR OFF GRID SENOIDAL PURA JFA 5000W 48/220V C/ GER RACK " and i['CONSTRUÇÃO DE CUSTO '] == "INVERSOR OFF GRID SENOIDAL PURA JFA 5000W 48/220V C/  GER RACK " and i['Anuncio'] == "Premium":
                    return round(i['PREÇO ML '], 2)

    def parse_location(self, response):
        
        name = response.meta['name']
        url = response.meta['url']
        new_price_float = response.meta['price']
        tipo = response.meta['tipo']
        cupom = response.meta['cupom']
        parcelado = self.get_price_previsto(tipo)
        loja = response.meta['loja']
        lugar = response.xpath('//*[@id="profile"]/div/div[2]/div[1]/div[3]/p/text()').get()
    
        if url not in urls_visitadas:
            urls_visitadas.append(url)
        else:
            return
        doc.add_paragraph(f'Modelo: {self.option_selected_new}')
        doc.add_paragraph(f'URL: {url}')
        doc.add_paragraph(f'Nome: {name}')
        doc.add_paragraph(f'Preço: {new_price_float}')
        doc.add_paragraph(f'Preço Previsto: {parcelado}')
        doc.add_paragraph(f'Loja: {loja}')
        doc.add_paragraph(f'Tipo: {tipo}')
        doc.add_paragraph(f'Lugar: {lugar}')
        doc.add_paragraph(f'Cupom: {cupom}')
        doc.add_paragraph("--------------------------------------------------------------------")
        doc.add_paragraph('')
        
        yield {
            'url': url,
            'name': name,
            'price': new_price_float,
            'price_previsto': parcelado,
            'loja': loja,
            'tipo': tipo,
            'lugar': lugar
        }
        if self.option_selected_new:
            doc.save(fr"dados/{self.option_selected_new.replace('/', '-')}.docx")
        else:
            doc.save("dados/documento_sem_nome.docx")

